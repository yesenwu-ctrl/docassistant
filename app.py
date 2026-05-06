from datetime import datetime
from io import BytesIO
import re

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from processor import ContentProcessor


st.set_page_config(page_title="AI擬答助理", layout="wide")

proc = ContentProcessor()


st.markdown(
    """
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Noto+Serif+TC:wght@600;700&family=Space+Grotesk:wght@500;700&display=swap');

        :root {
            --ink: #1e2019;
            --muted: #6a6a5f;
            --paper: #fffaf0;
            --sage: #dfe8cf;
            --clay: #c96f4a;
            --moss: #47533d;
            --line: rgba(30, 32, 25, 0.12);
        }

        .stApp {
            color: var(--ink);
            background:
                radial-gradient(circle at 8% 8%, rgba(201, 111, 74, 0.22), transparent 28rem),
                radial-gradient(circle at 90% 18%, rgba(111, 138, 86, 0.26), transparent 26rem),
                linear-gradient(135deg, #fffaf0 0%, #f4ecd9 48%, #eef1df 100%);
        }

        [data-testid="stSidebar"] {
            background: rgba(255, 250, 240, 0.86);
            border-right: 1px solid var(--line);
            backdrop-filter: blur(18px);
        }

        h1, h2, h3 {
            font-family: 'Noto Serif TC', serif;
            letter-spacing: -0.035em;
        }

        .hero {
            position: relative;
            overflow: hidden;
            padding: 2.4rem;
            border: 1px solid var(--line);
            border-radius: 34px;
            background:
                linear-gradient(140deg, rgba(255, 250, 240, 0.92), rgba(223, 232, 207, 0.72)),
                repeating-linear-gradient(90deg, rgba(71, 83, 61, 0.05) 0 1px, transparent 1px 18px);
            box-shadow: 0 28px 80px rgba(71, 83, 61, 0.14);
        }

        .hero:after {
            content: "";
            position: absolute;
            right: -4rem;
            top: -5rem;
            width: 16rem;
            height: 16rem;
            border-radius: 999px;
            background: rgba(201, 111, 74, 0.18);
        }

        .eyebrow {
            font-family: 'Space Grotesk', sans-serif;
            color: var(--clay);
            font-size: 0.82rem;
            font-weight: 700;
            letter-spacing: 0.12em;
            text-transform: uppercase;
        }

        .hero h1 {
            max-width: 820px;
            margin: 0.35rem 0 0.8rem;
            font-size: clamp(2.2rem, 6vw, 4.8rem);
            line-height: 0.96;
        }

        .hero p {
            max-width: 760px;
            color: var(--muted);
            font-size: 1.05rem;
            line-height: 1.8;
        }

        .metric-card {
            min-height: 118px;
            padding: 1.15rem;
            border-radius: 24px;
            border: 1px solid var(--line);
            background: rgba(255, 255, 255, 0.52);
        }

        .metric-card strong {
            display: block;
            color: var(--moss);
            font-family: 'Space Grotesk', sans-serif;
            font-size: 1.7rem;
        }

        .stButton > button {
            border: 0;
            border-radius: 999px;
            padding: 0.82rem 1.4rem;
            background: linear-gradient(135deg, var(--moss), #273022);
            color: #fffaf0;
            box-shadow: 0 16px 36px rgba(71, 83, 61, 0.24);
        }

        .stTextInput input, .stTextArea textarea, .stSelectbox [data-baseweb="select"] {
            border-radius: 18px;
        }

        [data-testid="stExpander"] {
            border: 1px solid var(--line);
            border-radius: 24px;
            background: rgba(255, 250, 240, 0.55);
        }
    </style>
    """,
    unsafe_allow_html=True,
)


def append_text_section(buffer, title, content):
    if content and content.strip():
        buffer.append(f"## {title}\n{content.strip()}")


def collect_uploaded_content(files, include_images=False):
    text_sections = []
    image_data_urls = []
    warnings = []

    for file in files or []:
        lower_name = file.name.lower()
        try:
            if lower_name.endswith(".docx"):
                append_text_section(text_sections, f"檔案：{file.name}", proc.read_docx(file))
            elif lower_name.endswith(".pdf"):
                append_text_section(text_sections, f"檔案：{file.name}", proc.read_pdf(file))
            elif lower_name.endswith((".png", ".jpg", ".jpeg", ".webp")):
                if include_images:
                    image_data_urls.append(proc.image_to_data_url(file))
                    append_text_section(text_sections, f"圖片：{file.name}", "已附上圖片，請直接閱讀圖片內容。")
                else:
                    warnings.append(f"已略過圖片「{file.name}」，因目前選擇純文字模式。")
            elif lower_name.endswith(".txt") or lower_name.endswith(".md"):
                append_text_section(text_sections, f"檔案：{file.name}", file.getvalue().decode("utf-8", errors="ignore"))
            else:
                warnings.append(f"暫不支援「{file.name}」的檔案格式。")
        except Exception:
            warnings.append(f"讀取「{file.name}」時發生錯誤，請改用貼上文字。")

    return "\n\n".join(text_sections), image_data_urls, warnings


AI_SERVICE_CONFIGS = {
    "OpenRouter": {
        "base_url": "https://openrouter.ai/api/v1",
        "models": {
            "DeepSeek Chat": "deepseek/deepseek-chat",
            "DeepSeek R1": "deepseek/deepseek-r1",
        },
        "caption": "使用 OpenRouter 的文字模型",
        "placeholder": "貼上 OpenRouter API Key",
    },
    "Gemini": {
        "base_url": "https://generativelanguage.googleapis.com/v1beta/openai",
        "models": {
            "Gemini 2.5 Flash": "gemini-2.5-flash",
            "Gemini 2.5 Pro": "gemini-2.5-pro",
            "Gemini 2.5 Flash-Lite": "gemini-2.5-flash-lite",
            "Gemini 2.0 Flash": "gemini-2.0-flash",
            "Gemini 2.0 Flash-Lite": "gemini-2.0-flash-lite",
        },
        "caption": "使用 Google Gemini 文字模型",
        "placeholder": "貼上 Gemini API Key",
    },
}


def resolve_ai_config(ai_service):
    return AI_SERVICE_CONFIGS[ai_service]


def clean_markdown_text(text):
    cleaned = re.sub(r"^\s{0,3}#{1,6}\s*", "", text.strip())
    cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"__(.*?)__", r"\1", cleaned)
    return cleaned.strip()


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Title", 22, RGBColor(47, 58, 43)),
        ("Heading 1", 16, RGBColor(47, 58, 43)),
        ("Heading 2", 13, RGBColor(73, 83, 61)),
        ("Heading 3", 12, RGBColor(94, 76, 58)),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft JhengHei"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_meta_table(doc, response_format, model):
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    entries = [
        ("輸出格式", response_format),
        ("使用模型", model),
        ("產製時間", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("用途", "AI 擬答草稿"),
    ]
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            label, value = entries[row_index * 2 + col_index]
            cell.text = f"{label}\n{value}"
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Microsoft JhengHei"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(74, 83, 62)


def add_response_to_docx(doc, text):
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", line)
        question_match = re.match(r"^(Q\d+[\.\uff0e\uff1a:、\s].+)$", line, re.IGNORECASE)
        chinese_section_match = re.match(r"^([一二三四五六七八九十]+、.+)$", line)
        bullet_match = re.match(r"^[-*•]\s+(.+)$", line)
        number_match = re.match(r"^\d+[.)、]\s+(.+)$", line)
        bold_heading_match = re.match(r"^\*\*(.+)\*\*$", line)

        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            doc.add_heading(clean_markdown_text(heading_match.group(2)), level=level)
        elif question_match:
            doc.add_heading(clean_markdown_text(question_match.group(1)), level=1)
        elif chinese_section_match and len(line) <= 42:
            doc.add_heading(clean_markdown_text(chinese_section_match.group(1)), level=2)
        elif bold_heading_match and len(line) <= 60:
            doc.add_heading(clean_markdown_text(bold_heading_match.group(1)), level=3)
        elif bullet_match:
            doc.add_paragraph(clean_markdown_text(bullet_match.group(1)), style="List Bullet")
        elif number_match:
            doc.add_paragraph(clean_markdown_text(number_match.group(1)), style="List Number")
        else:
            doc.add_paragraph(clean_markdown_text(line))


def build_docx_bytes(text, response_format, model):
    doc = Document()
    style_document(doc)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("AI擬答草稿")
    add_meta_table(doc, response_format, model)
    doc.add_paragraph("")
    add_response_to_docx(doc, text)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def build_prompt(reply_goal, reply_tone, response_format, source_content, reference_content, user_direction):
    format_instruction = f"輸出格式使用「{response_format}」。"
    if response_format == "採訪稿":
        format_instruction = """
輸出格式使用「採訪稿」，並嚴格依照以下架構：
1. 先用一段話從整體面向概述現況，交代背景、問題脈絡與目前推動方向。
2. 接著分點說明，每一點都必須有清楚小標題，內容聚焦作法、政策效益或已呈現成果。
3. 最後歸納重點，針對問題寫一段簡短回應。這段不用強調數據，重點放在如何解決問題、回應需求或改善現況。
若使用者指定多個題號，請每一題都各自套用上述架構。
""".strip()

    system_prompt = f"""
你是一位謹慎、精準且善於擬稿的中文助理。你的任務不是只整理摘要，而是根據使用者提供的來源內容與參考資料，直接產出可使用的回覆草稿。

寫作規則：
1. 優先回答對方真正需要的內容，不要只做條列摘錄。
2. 若來源內容資訊不足，請在回覆中以溫和方式標明「需要補充確認」的地方，不要捏造。
3. 參考資料只能作為補強脈絡，不得覆蓋來源內容中的明確事實。
4. 語氣使用「{reply_tone}」。
5. {format_instruction}
6. 若使用者指定題號或段落，例如 Q1、Q4，必須依照指定題號分段回覆，保留題號作為小標題；若使用者說「只要」或「僅」回覆特定題號，就不要補寫其他題號。
7. 若來源內容本身包含訪綱題號，也要逐題回覆，不要自行改寫成單一總論。
8. 避免使用自稱 AI、要求人工檢查、或過度說明資料來源的制式語句；若資訊不足，改用自然文字說明「仍有待確認」或「建議補充確認」。
9. 若內容涉及法規、承諾、數字或時程，請用保守措辭，避免做超出資料範圍的承諾。
""".strip()

    user_prompt = f"""
# 我想請 AI 完成的回覆任務
{reply_goal.strip() or "請根據來源內容擬一份完整、清楚、可直接修改使用的回覆。"}

# 使用者額外要求（最高優先）
{user_direction.strip() or "無"}

# 需要回覆的來源內容
{source_content.strip()}

# 可參考的背景資料
{reference_content.strip() or "無"}
""".strip()

    return system_prompt, user_prompt


with st.sidebar:
    st.caption("AI SERVICE")
    st.header("AI 服務與金鑰")
    ai_service = st.selectbox("AI 服務", list(AI_SERVICE_CONFIGS.keys()))
    ai_config = resolve_ai_config(ai_service)
    model_label = st.selectbox("模型", list(ai_config["models"].keys()))
    api_key = st.text_input("API Key", type="password", placeholder=ai_config["placeholder"])
    base_url = ai_config["base_url"]
    model = ai_config["models"][model_label]
    st.caption(ai_config["caption"])
    if ai_service == "Gemini":
        st.caption("若免費額度或速率達上限，系統會顯示額度上限提示；可稍後再試或改用較輕量模型。")

    st.divider()
    st.caption("WRITING CONTROL")
    reply_tone = st.selectbox("回覆語氣", ["正式清楚", "友善專業", "簡潔直接", "溫和說明", "政策幕僚風格"])
    response_format = st.selectbox("輸出格式", ["採訪稿", "完整回覆信件", "正式公文式回覆", "條列說明", "Line/簡訊短回覆", "Markdown 草稿"])
    user_direction = st.text_area("特別注意事項", placeholder="例如：避免承諾時程、引用資安法規、需要先感謝對方...")


st.markdown(
    """
    <section class="hero">
        <div class="eyebrow">AI Reply Assistant</div>
        <h1>AI擬答助理</h1>
        <p>
            上傳來文或貼上文字，再補充參考資料與回覆方向，
            讓 AI 直接擬好一份可修改、可送出的草稿。
        </p>
    </section>
    """,
    unsafe_allow_html=True,
)

metric_cols = st.columns(3)
with metric_cols[0]:
    st.markdown('<div class="metric-card"><strong>01</strong>貼上 API Key</div>', unsafe_allow_html=True)
with metric_cols[1]:
    st.markdown('<div class="metric-card"><strong>02</strong>放入來文與參考資料</div>', unsafe_allow_html=True)
with metric_cols[2]:
    st.markdown('<div class="metric-card"><strong>03</strong>生成可直接修改的回覆</div>', unsafe_allow_html=True)

st.write("")

left, right = st.columns([1.05, 0.95], gap="large")

with left:
    st.subheader("需要回覆的內容")
    source_tabs = st.tabs(["檔案上傳", "網頁連結", "文字貼上"])

    with source_tabs[0]:
        uploaded_files = st.file_uploader(
            "上傳來文、PDF、Word 或文字檔",
            type=["docx", "pdf", "txt", "md"],
            accept_multiple_files=True,
        )

    with source_tabs[1]:
        url_input = st.text_input("貼上公開網頁或 Google 表單連結")

    with source_tabs[2]:
        raw_text = st.text_area("直接貼上需要回覆的內容", height=220, placeholder="把對方來信、問題、會議紀錄或申請內容貼在這裡。")

with right:
    st.subheader("參考資料與任務")
    reply_goal = st.text_area(
        "你希望 AI 幫你怎麼回？",
        height=120,
        placeholder="例如：請幫我回覆承辦單位，說明本案可先提供初步資料，但正式時程需待內部確認。",
    )
    reference_text = st.text_area(
        "可參考資料",
        height=220,
        placeholder="貼上政策背景、過往回覆、內部原則、注意事項或希望 AI 參考但不要逐字照抄的資料。",
    )
    reference_files = st.file_uploader(
        "補充參考檔案",
        type=["docx", "pdf", "txt", "md"],
        accept_multiple_files=True,
        key="reference_files",
    )

generate = st.button("生成 AI 回覆草稿", type="primary", use_container_width=True)

if generate:
    source_parts = []
    reference_parts = []
    image_data_urls = []
    warnings = []

    uploaded_text, uploaded_images, uploaded_warnings = collect_uploaded_content(
        uploaded_files,
        include_images=False,
    )
    append_text_section(source_parts, "上傳內容", uploaded_text)
    image_data_urls.extend(uploaded_images)
    warnings.extend(uploaded_warnings)

    if url_input:
        url_text, url_error = proc.fetch_url(url_input)
        if url_error:
            warnings.append(url_error)
        else:
            append_text_section(source_parts, "網頁內容", url_text)

    append_text_section(source_parts, "貼上文字", raw_text)

    reference_file_text, _, reference_warnings = collect_uploaded_content(reference_files, include_images=False)
    append_text_section(reference_parts, "貼上參考資料", reference_text)
    append_text_section(reference_parts, "參考檔案", reference_file_text)
    warnings.extend(reference_warnings)

    for warning in warnings:
        st.warning(warning)

    source_content = "\n\n".join(source_parts)
    reference_content = "\n\n".join(reference_parts)

    if not source_content.strip():
        st.error("請至少提供一段需要回覆的內容。")
    elif not api_key.strip():
        st.error("請先在左側輸入 API Key。")
    else:
        system_prompt, user_prompt = build_prompt(
            reply_goal,
            reply_tone,
            response_format,
            source_content,
            reference_content,
            user_direction,
        )

        with st.spinner("AI 正在閱讀資料並擬回覆..."):
            result, ai_error = proc.call_ai(
                api_key,
                base_url,
                model,
                system_prompt,
                user_prompt,
                image_data_urls=image_data_urls,
            )

        if ai_error:
            st.error(ai_error)
        else:
            st.success("已生成回覆草稿，可再依實際情況微調後使用。")
            st.markdown("### AI 回覆草稿")
            st.markdown(result)
            docx_bytes = build_docx_bytes(result, response_format, model)
            st.download_button(
                "下載 Word 檔",
                data=docx_bytes,
                file_name=f"AI擬答草稿_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

            with st.expander("查看本次送給 AI 的整理內容"):
                st.markdown(user_prompt)
